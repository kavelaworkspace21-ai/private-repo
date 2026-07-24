"""
Module 2 — AI Drafting Engine
Structured document drafting via SSE streaming.
8 core templates: Legal Notice, Cheque Dishonour, Bail (Anticipatory + Regular),
Affidavit, RTI Application, Consumer Complaint, Vakalatnama.
"""
import os
import io
import json
import re
from app.util.time import utcnow
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field, field_validator
from typing import Any

from sqlalchemy.orm import Session

from app.db.session import get_db
from app.auth.dependencies import get_current_user, require_ai_access
from app.services.ratelimit import ai_limiter
from app.services.tenancy import write_audit
from app.models.user import User

router = APIRouter()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")

# ── Document type catalogue ───────────────────────────────────────────────────
DOCUMENT_TYPES: dict[str, dict] = {
    "custom_document": {
        "label": "Custom Document (describe what you need)",
        "fields": ["document_title", "instructions", "parties_and_facts"],
        "system_prompt": """You are an expert Indian legal drafter. Draft the document the advocate
describes, in correct Indian legal form and language, ready for the advocate to review and finalise.
Use proper structure appropriate to the document type (title, parties, recitals/body, prayer or relief,
and a signature / verification block where relevant). Cite only provisions you are confident apply, and
leave [●] placeholders for any unknown fact rather than inventing it.""",
        "user_prompt_template": """Draft the following document for an Indian legal matter:

Document: {document_title}

Instructions / what it must achieve:
{instructions}

Parties / facts / details:
{parties_and_facts}

Produce the complete, formatted document, ready for advocate review.""",
    },
    "crpc_application": {
        "label": "Criminal Misc. Application (any BNSS/CrPC section)",
        "fields": ["court_name", "case_number", "applicant_name", "applicant_role",
                   "opposite_party", "provision", "purpose", "facts", "grounds", "relief_sought"],
        "system_prompt": """You are an expert Indian criminal-side drafter. Draft a miscellaneous
application for the provision the advocate names — ANY section of the Bharatiya Nagarik Suraksha
Sanhita, 2023 (or CrPC, 1973 for pre-1 July 2024 matters): e.g. recall of witness, further
investigation, release of property, modification of bail conditions, condonation of delay, exemption
from appearance, discharge, etc.
REQUIREMENTS: standard misc-application structure (cause title with the pending case number,
provision heading, showeth paragraphs — pendency, facts, how each ingredient of the provision is
satisfied, grounds — then a specific prayer tracking the provision's language). If RETRIEVED
PROVISIONS include the section's verbatim text, mirror its language when stating what the court is
empowered to do; NEVER paraphrase statute from memory. Where both codes could apply, use the BNSS
section and note the CrPC equivalent for pre-2024 matters. Keep [●] for unknown facts.""",
        "user_prompt_template": """Draft a criminal miscellaneous application:

Court: {court_name}
Pending case: {case_number}
Applicant: {applicant_name} ({applicant_role})
Opposite party: {opposite_party}
Provision invoked: {provision}
Purpose of the application: {purpose}
Relevant facts: {facts}
Grounds: {grounds}
Relief sought: {relief_sought}

Produce the complete, formatted application, ready for advocate review.""",
    },
    "cpc_application": {
        "label": "Civil Misc. Application (any CPC Order/Section)",
        "fields": ["court_name", "case_number", "applicant_name", "applicant_role",
                   "opposite_party", "provision", "purpose", "facts", "grounds", "relief_sought"],
        "system_prompt": """You are an expert Indian civil-side drafter. Draft an interlocutory /
miscellaneous application for the provision the advocate names — ANY Order/Rule or Section of the
Code of Civil Procedure, 1908: e.g. O.VI R.17 amendment, O.IX R.13 set-aside, O.XXVI commissioner,
O.XXXVIII R.5 attachment, O.XL receiver, s.148 enlargement of time, s.152 correction, s.151 inherent
powers, etc.
REQUIREMENTS: standard I.A. structure (cause title in the pending suit, provision heading — read
with s.151 CPC where appropriate — showeth paragraphs: pendency and stage, facts necessitating the
application, how the provision's requirements are met, grounds; then a specific prayer). If RETRIEVED
PROVISIONS include the section's verbatim text, track its language; NEVER quote CPC Orders/Rules from
memory — state their effect generally and keep the citation to the number the advocate gave. Note
where an affidavit in support is required. Keep [●] for unknown facts.""",
        "user_prompt_template": """Draft a civil miscellaneous (interlocutory) application:

Court: {court_name}
Pending suit/proceeding: {case_number}
Applicant: {applicant_name} ({applicant_role})
Opposite party: {opposite_party}
Provision invoked: {provision}
Purpose of the application: {purpose}
Relevant facts: {facts}
Grounds: {grounds}
Relief sought: {relief_sought}

Produce the complete, formatted application, ready for advocate review.""",
    },
    "legal_notice_cpc": {
        "label": "Legal Notice (s.80 CPC)",
        "fields": ["sender_name", "sender_address", "recipient_name", "recipient_address",
                   "subject_matter", "amount_or_relief", "cause_of_action", "demand_period_days"],
        "system_prompt": """You are an expert Indian legal drafter. Draft a formal legal notice under Section 80 of the Code of Civil Procedure, 1908 (CPC).

REQUIREMENTS:
- Full professional format: date, sender address, recipient address, subject line, body, demand, consequence of non-compliance, advocate's signature block
- Cite Section 80 CPC correctly (pre-suit notice before filing suit against Government) or Section 80A where applicable
- State facts clearly: cause of action, amounts/reliefs claimed, timeline
- Include a specific demand period (usually 60 days for Govt, or as instructed)
- End with consequences: filing of civil suit for recovery + costs + interest
- Use formal legal language appropriate for Indian courts
- Format as a complete, ready-to-send document""",
        "user_prompt_template": """Draft a legal notice under Section 80 CPC with the following details:

Sender: {sender_name}
Sender's Address: {sender_address}

Recipient: {recipient_name}
Recipient's Address: {recipient_address}

Subject / Matter: {subject_matter}
Amount / Relief Sought: {amount_or_relief}
Cause of Action: {cause_of_action}
Notice Period: {demand_period_days} days

Produce the complete, formatted legal notice document.""",
    },

    "cheque_dishonour": {
        "label": "Cheque Dishonour Notice (s.138 NI Act)",
        "fields": ["drawer_name", "drawer_address", "payee_name", "payee_address",
                   "cheque_number", "cheque_date", "cheque_amount", "bank_name",
                   "dishonour_date", "dishonour_reason", "underlying_liability"],
        "system_prompt": """You are an expert Indian legal drafter specialising in Negotiable Instruments Act, 1881 matters.

Draft a statutory demand notice under Section 138 read with Section 142 of the Negotiable Instruments Act, 1881.

REQUIREMENTS:
- Notice must be sent within 30 days of receiving dishonour memo
- Demand payment within 15 days of receipt of notice (as required by proviso to s.138)
- State all essential facts: cheque details, dishonour particulars, underlying legally enforceable debt/liability
- Warn that failure to pay within 15 days will result in criminal complaint u/s 138 NI Act (imprisonment up to 2 years / fine up to twice the cheque amount)
- Reference the banking dishonour memo
- Use correct legal terminology
- Format as a complete document ready for dispatch (by speed post / courier)""",
        "user_prompt_template": """Draft a legal notice under Section 138 of the Negotiable Instruments Act, 1881 with these details:

Drawer (Accused): {drawer_name}, {drawer_address}
Payee (Complainant): {payee_name}, {payee_address}

Cheque No.: {cheque_number}
Cheque Date: {cheque_date}
Cheque Amount: ₹{cheque_amount}
Drawn on Bank: {bank_name}
Date of Dishonour: {dishonour_date}
Reason for Dishonour: {dishonour_reason}
Underlying Liability/Debt: {underlying_liability}

Produce the complete statutory demand notice.""",
    },

    "anticipatory_bail": {
        "label": "Anticipatory Bail Application (s.482 BNSS / s.438 CrPC)",
        "fields": ["applicant_name", "applicant_address", "fir_number", "police_station",
                   "offences_alleged", "brief_facts", "grounds_for_bail", "court_name"],
        "system_prompt": """You are an expert Indian criminal law drafter.

Draft a formal Application for Anticipatory Bail under Section 482 of the Bharatiya Nagarik Suraksha Sanhita, 2023 (BNSS) [formerly Section 438 of CrPC, 1973] before the High Court / Sessions Court.

REQUIREMENTS:
- Proper case title: "In the matter of anticipatory bail application..."
- List applicant details, FIR particulars, offences alleged
- State grounds for anticipatory bail clearly — cite relevant Supreme Court judgments (Gurbaksh Singh Sibbia v. State of Punjab, Siddharam Satlingappa v. State of Maharashtra, etc.)
- Address the factors courts consider: nature of accusation, antecedents, possibility of fleeing justice, likelihood of repeating offence, custodial interrogation necessity
- Prayer clause: specific directions sought
- Undertakings: availability for interrogation, surrendering passport, not influencing witnesses
- Advocate verification / certification
- Format as a complete, court-ready document with cause title, synopsis, list of dates, grounds, and prayer""",
        "user_prompt_template": """Draft an Anticipatory Bail Application under Section 482 BNSS (s.438 CrPC) with these details:

Court: {court_name}
Applicant: {applicant_name}, {applicant_address}
FIR No.: {fir_number}, Police Station: {police_station}
Offences Alleged: {offences_alleged}
Brief Facts of the Case: {brief_facts}
Grounds for Bail: {grounds_for_bail}

Produce the complete application with cause title, synopsis, grounds, prayer, and undertakings.""",
    },

    "regular_bail": {
        "label": "Bail Application (s.480 BNSS / s.437 CrPC)",
        "fields": ["applicant_name", "applicant_address", "case_number", "police_station",
                   "date_of_arrest", "offences_alleged", "brief_facts", "grounds_for_bail",
                   "court_name", "days_in_custody"],
        "system_prompt": """You are an expert Indian criminal law drafter.

Draft a formal Bail Application under Section 480 of the Bharatiya Nagarik Suraksha Sanhita, 2023 (BNSS) [formerly Section 437/439 CrPC] before the competent court.

REQUIREMENTS:
- Complete cause title with case/FIR reference
- State: accused details, date of arrest, custody duration, offences charged
- Grounds for bail: presumption of innocence, no criminal antecedents, roots in community, cooperation with investigation, health/family grounds as applicable
- Address triple test: flight risk, tampering with evidence, threat to witnesses
- Cite key SC judgments: Arnesh Kumar v. State of Bihar, Sanjay Chandra v. CBI, Dataram Singh v. State of U.P.
- Proposed conditions if bail is granted
- Prayer clause
- Complete court-ready format with verification""",
        "user_prompt_template": """Draft a Bail Application under Section 480 BNSS with these details:

Court: {court_name}
Applicant/Accused: {applicant_name}, {applicant_address}
Case/FIR No.: {case_number}, Police Station: {police_station}
Date of Arrest: {date_of_arrest}
Days in Custody: {days_in_custody}
Offences: {offences_alleged}
Brief Facts: {brief_facts}
Grounds for Bail: {grounds_for_bail}

Produce the complete bail application with cause title, grounds, prayer, and verification.""",
    },

    "affidavit": {
        "label": "Affidavit / Sworn Statement",
        "fields": ["deponent_name", "deponent_age", "deponent_address", "deponent_occupation",
                   "purpose_of_affidavit", "facts_to_state", "place_of_swearing"],
        "system_prompt": """You are an expert Indian legal drafter.

Draft a formal Affidavit / Sworn Statement compliant with Indian legal requirements (Code of Civil Procedure, 1908, Order XIX; Oaths Act, 1969).

REQUIREMENTS:
- Standard format: deponent's full name, age, address, occupation
- "I, [name], son/daughter of [father's name], aged [x] years, residing at [address], do hereby solemnly affirm and state as under:"
- Numbered paragraphs for each fact
- "I say that..." / "I state that..." / "I further say that..." construction
- Verification clause: "Verified at [place] on this [date] day of [month], [year] that the contents of paragraphs [x] to [y] are true to my knowledge and the rest are true to my belief and information."
- Signature/thumb impression line
- Space for Notary/Oath Commissioner attestation
- If for court: reference to the relevant proceeding""",
        "user_prompt_template": """Draft an Affidavit with the following details:

Deponent: {deponent_name}, Age: {deponent_age}
Address: {deponent_address}
Occupation: {deponent_occupation}
Purpose of Affidavit: {purpose_of_affidavit}
Facts to State: {facts_to_state}
Place of Swearing: {place_of_swearing}

Produce the complete affidavit with verification clause.""",
    },

    "rti_application": {
        "label": "RTI Application (RTI Act 2005)",
        "fields": ["applicant_name", "applicant_address", "applicant_contact",
                   "public_authority", "pio_designation", "information_sought",
                   "period_of_information", "purpose"],
        "system_prompt": """You are an expert in the Right to Information Act, 2005.

Draft a formal RTI Application under Section 6(1) of the Right to Information Act, 2005 to the Public Information Officer (PIO).

REQUIREMENTS:
- Address to: The Public Information Officer (name of public authority)
- Subject: Application under Section 6(1) of the Right to Information Act, 2005
- Applicant details: name, address, contact
- Specific, precise information sought (vague requests get rejected)
- Reference to relevant period/files
- Statement that the applicant is an Indian citizen
- Application fee note: ₹10 (Central Govt) / as applicable to State
- Time limit reminder: information to be provided within 30 days (u/s 7)
- Mention right to first appeal u/s 19 if not responded to within 30 days
- Date, place, signature
- Format as complete ready-to-submit application""",
        "user_prompt_template": """Draft an RTI Application with these details:

Applicant: {applicant_name}
Address: {applicant_address}
Contact: {applicant_contact}

To: The Public Information Officer
Public Authority: {public_authority}
Designation of PIO: {pio_designation}

Information Sought: {information_sought}
Period of Information: {period_of_information}
Purpose/Context: {purpose}

Produce the complete RTI application.""",
    },

    "consumer_complaint": {
        "label": "Consumer Complaint (Consumer Protection Act 2019)",
        "fields": ["complainant_name", "complainant_address", "opposite_party_name",
                   "opposite_party_address", "product_or_service", "date_of_purchase",
                   "amount_paid", "deficiency_or_defect", "relief_sought", "forum"],
        "system_prompt": """You are an expert in Indian consumer law, specifically the Consumer Protection Act, 2019 and the Consumer Protection (Consumer Disputes Redressal Commissions) Rules, 2020.

Draft a formal Consumer Complaint to the appropriate District/State/National Consumer Disputes Redressal Commission.

REQUIREMENTS:
- Correct forum based on pecuniary jurisdiction (District: up to ₹50 lakhs; State: ₹50 lakhs–₹2 crores; National: above ₹2 crores)
- Cause title: Consumer Complaint No. ___ / 20__
- Complainant and Opposite Party details
- Facts: purchase/service details, deficiency of service / defect in goods, date of cause of action
- Legal provisions: Section 2(7) [consumer definition], Section 2(11) [deficiency], Section 2(10) [defect], Section 35/47/58 [filing jurisdiction]
- Relief sought: replacement/refund/compensation/punitive damages/costs
- Annexures list (purchase receipt, warranty, correspondence, etc.)
- Verification and affidavit in support
- Complete ready-to-file format""",
        "user_prompt_template": """Draft a Consumer Complaint with these details:

Forum: {forum}
Complainant: {complainant_name}, {complainant_address}
Opposite Party: {opposite_party_name}, {opposite_party_address}

Product/Service: {product_or_service}
Date of Purchase/Transaction: {date_of_purchase}
Amount Paid: ₹{amount_paid}
Deficiency/Defect: {deficiency_or_defect}
Relief Sought: {relief_sought}

Produce the complete consumer complaint with cause title, facts, legal grounds, prayer, and verification.""",
    },

    "vakalatnama": {
        "label": "Vakalatnama / Memo of Appearance",
        "fields": ["client_name", "client_address", "advocate_name", "advocate_enrollment",
                   "bar_council", "case_title", "case_number", "court_name",
                   "party_capacity"],
        "system_prompt": """You are an expert Indian legal drafter.

Draft a formal Vakalatnama (Power of Attorney / Memo of Appearance) authorising an advocate to appear on behalf of a party in court proceedings.

REQUIREMENTS:
- Standard Vakalatnama format used across Indian courts
- Client (executant) authorises named advocate (and their associates/juniors)
- Powers granted: appear, act, plead, file documents, receive notices, accept service of process, sign pleadings, compromise/settle (if applicable), file appeals, engage other advocates
- Limitations if any
- Enrollment details of advocate
- Proper execution clause: signed by client, witnessed
- Stamp paper requirement note (vary by state — typically ₹100 stamp paper for High Courts)
- Separate section for advocate's undertaking / acceptance
- Court-ready format""",
        "user_prompt_template": """Draft a Vakalatnama with these details:

Court: {court_name}
Case Title: {case_title}
Case Number: {case_number}
Client (Executant): {client_name}, {client_address}
Party Capacity: {party_capacity} (Plaintiff/Defendant/Petitioner/Respondent/Appellant)

Advocate Authorised: {advocate_name}
Enrollment No.: {advocate_enrollment}
Bar Council: {bar_council}

Produce the complete Vakalatnama with execution clause and advocate's acceptance.""",
    },

    "writ_petition": {
        "label": "Writ Petition (Art. 226 / 32)",
        "fields": ["petitioner_name", "petitioner_address", "respondent_authority",
                   "court_name", "rights_violated", "brief_facts", "relief_sought"],
        "system_prompt": """You are an expert Indian constitutional law drafter.

Draft a Writ Petition under Article 226 of the Constitution of India (before a High Court) or Article 32 (before the Supreme Court), as indicated by the court named.

REQUIREMENTS:
- Correct cause title and jurisdiction clause (Art. 226 for High Court / Art. 32 for Supreme Court)
- Parties: petitioner(s) and respondent authority/State
- Synopsis and list of dates
- Facts: concise, numbered
- Grounds: the legal/fundamental rights violated (cite the specific Articles and any statute), and why the impugned action is illegal/arbitrary/ultra vires
- Nature of writ sought (mandamus / certiorari / prohibition / quo warranto / habeas corpus) with reasons
- Prayer clause with specific reliefs and interim relief if applicable
- Affidavit / verification in support
- Court-ready format. Leave [●] placeholders where a fact is unknown — never invent facts, dates, or authorities.""",
        "user_prompt_template": """Draft a Writ Petition with these details:

Court: {court_name}
Petitioner: {petitioner_name}, {petitioner_address}
Respondent Authority: {respondent_authority}
Rights / Provisions Violated: {rights_violated}
Brief Facts: {brief_facts}
Relief Sought: {relief_sought}

Produce the complete writ petition with cause title, synopsis, facts, grounds, prayer, and verification.""",
    },

    "divorce_petition": {
        "label": "Divorce Petition",
        "fields": ["petitioner_name", "respondent_name", "marriage_date", "place_of_marriage",
                   "applicable_law", "ground_for_divorce", "court_name", "brief_facts", "relief_sought"],
        "system_prompt": """You are an expert Indian matrimonial law drafter.

Draft a Petition for Divorce before the Family Court / District Court, under the personal law indicated by the user (e.g. Section 13 of the Hindu Marriage Act, 1955; the Special Marriage Act, 1954; the Divorce Act, 1869; or Muslim personal law as applicable).

REQUIREMENTS:
- Cause title and jurisdiction clause (where the marriage was solemnised / parties last resided together / respondent resides)
- Parties' details and particulars of the marriage (date, place, status)
- Facts: numbered, concise
- Ground(s) for divorce stated with reference to the SPECIFIC section of the applicable Act — use ONLY the law the user names; if unclear, state that the advocate must confirm the applicable personal law and jurisdiction
- Statement on no collusion/condonation; limitation/maintainability
- Prayer clause (dissolution of marriage and any ancillary relief sought)
- Verification and affidavit
- Sensitive, factual, non-inflammatory language. Leave [●] placeholders rather than inventing facts.""",
        "user_prompt_template": """Draft a Divorce Petition with these details:

Court: {court_name}
Applicable Law: {applicable_law}
Petitioner: {petitioner_name}
Respondent: {respondent_name}
Marriage Date: {marriage_date}, Place: {place_of_marriage}
Ground for Divorce: {ground_for_divorce}
Brief Facts: {brief_facts}
Relief Sought: {relief_sought}

Produce the complete divorce petition with cause title, facts, statutory ground, prayer, and verification.""",
    },
}

# ── Schemas ────────────────────────────────────────────────────────────────────
class DraftRequest(BaseModel):
    document_type: str
    fields: dict[str, Any]

    @field_validator("fields")
    @classmethod
    def _bounded_fields(cls, v: dict) -> dict:
        # Abuse guard: bounded form, not a free channel into the LLM prompt.
        if len(v) > 40:
            raise ValueError("Too many fields (max 40).")
        for k, val in v.items():
            if len(str(k)) > 64:
                raise ValueError("Field name too long (max 64 chars).")
            if len(str(val)) > 8000:
                raise ValueError(f"Field '{k}' too long (max 8000 chars).")
        return v


# ── WB-07: in-app editor actions + review-own-draft ──────────────────────────
EDIT_ACTIONS = {
    "expand":  "Expand the SELECTION with fuller drafting, keeping every fact and citation intact.",
    "shorten": "Tighten the SELECTION; remove redundancy; keep every fact, date and citation.",
    "reword":  "Reword the SELECTION in clearer professional drafting; same meaning, facts, citations.",
    "tone":    "Rewrite the SELECTION in the requested tone (see instruction), preserving substance.",
    "add_clause": "Draft the clause described in the instruction so it fits the SELECTION's context.",
}


class EditRequest(BaseModel):
    action: str = Field(pattern=r"^(expand|shorten|reword|tone|add_clause)$")
    selection: str = Field(min_length=1, max_length=20_000)
    context: str = Field(default="", max_length=40_000)     # surrounding text (optional)
    instruction: str = Field(default="", max_length=1_000)  # e.g. tone, clause description


@router.post("/edit", dependencies=[Depends(ai_limiter)])
def edit_selection(body: EditRequest, user: User = Depends(require_ai_access),
                   db: Session = Depends(get_db)):
    """Editor actions (pack WB-07): transform ONLY the selection. Pure transform — the
    caller applies it; saving still goes through /api/drafts (review status + versions).
    Never invents facts; [●] placeholders and citations must survive."""
    from app.ai.llm_config import ai_config
    cfg = ai_config()
    if not cfg["api_key"]:
        raise HTTPException(503, "The AI engine isn't connected (no model key on this server).")
    from app.ai.safety import sanitize_answer
    system = ("You are an Indian legal drafting editor. Apply ONE edit to the SELECTION and "
              "return ONLY the replacement text — no commentary, no markdown fences.\n"
              f"EDIT: {EDIT_ACTIONS[body.action]}\n"
              "RULES: never invent facts, names, dates or amounts; keep [●] placeholders; keep "
              "every statutory citation exactly; no outcome promises.")
    user_msg = ((f"INSTRUCTION: {body.instruction}\n\n" if body.instruction else "")
                + (f"SURROUNDING CONTEXT (do not return this):\n{body.context[:6000]}\n\n"
                   if body.context else "")
                + f"SELECTION:\n{body.selection}")
    from openai import OpenAI
    try:
        client = OpenAI(api_key=cfg["api_key"], base_url=cfg["base_url"], timeout=90, max_retries=0)
        resp = client.chat.completions.create(model=cfg["model"], max_tokens=2048, temperature=0.2,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user_msg}])
        out = sanitize_answer((resp.choices[0].message.content or "").strip())
    except Exception as e:
        raise HTTPException(502, f"Edit failed ({str(e)[:120]}) — nothing was changed; retry.")
    write_audit(db, tenant_id=user.tenant_id, user_id=user.id, action="draft_edit_action",
                entity="DraftEdit", detail=body.action)
    return {"replacement": out, "action": body.action}


class ReviewOwnDraft(BaseModel):
    content: str = Field(min_length=50, max_length=300_000)


@router.post("/review-draft", dependencies=[Depends(ai_limiter)])
def review_own_draft(body: ReviewOwnDraft, user: User = Depends(require_ai_access),
                     db: Session = Depends(get_db)):
    """Upload-your-own-draft Review mode (pack WB-07): ADVISORY, cited flags — issues,
    missing averments, limitation/jurisdiction — grounded on retrieved statute only.
    Metered as one research unit; a grounding failure refuses rather than guesses."""
    from app.models.billing import KIND_RESEARCH
    from app.services.entitlements import enforce_quota, meter
    enforce_quota(db, user, KIND_RESEARCH)
    from app.ai import rag
    from app.ai.safety import extract_citations, sanitize_answer, DRAFT_DISCLAIMER
    grounding = ""
    try:
        grounding = (rag.retrieve_by_section(body.content[:400]) or "") or rag.retrieve(body.content[:400])
    except Exception:
        pass
    from app.ai.llm_config import ai_config
    cfg = ai_config()
    if not cfg["api_key"]:
        raise HTTPException(503, "The AI engine isn't connected (no model key on this server).")
    system = ("You are a senior Indian advocate reviewing a junior's draft. Output markdown with "
              "EXACTLY these '## ' sections: Issues Found · Missing Averments · Limitation Flags · "
              "Jurisdiction Flags · Suggested Fixes.\n"
              "RULES: ADVISORY tone only; every legal proposition cites ONLY provisions in the "
              "RETRIEVED LAW block ('Section N of the <Act>'); if the retrieved law doesn't cover a "
              "point, say so; never predict outcomes; end with: \"" + DRAFT_DISCLAIMER + "\"\n\n"
              f"RETRIEVED LAW:\n{grounding or '(nothing retrieved)'}")
    from openai import OpenAI
    try:
        client = OpenAI(api_key=cfg["api_key"], base_url=cfg["base_url"], timeout=120, max_retries=0)
        resp = client.chat.completions.create(model=cfg["model"], max_tokens=2048, temperature=0.1,
            messages=[{"role": "system", "content": system},
                      {"role": "user", "content": f"DRAFT UNDER REVIEW:\n{body.content[:30000]}"}])
        review = sanitize_answer((resp.choices[0].message.content or "").strip())
    except Exception as e:
        raise HTTPException(502, f"Review failed ({str(e)[:120]}) — nothing was charged; retry.")
    verified = sorted(set(extract_citations(review)) & set(extract_citations(grounding)))
    meter(db, user, KIND_RESEARCH)
    write_audit(db, tenant_id=user.tenant_id, user_id=user.id, action="draft_review_mode",
                entity="DraftReview", detail=f"citations={len(verified)}")
    return {"review": review, "verified_citations": verified,
            "status": "ADVISORY_FOR_ADVOCATE", "disclaimer": DRAFT_DISCLAIMER}


class ExportRequest(BaseModel):
    document_type: str
    # 300k chars ≈ a very long draft; beyond that is not a document, it's abuse of the
    # DOCX/PDF renderers (CPU-bound on the small box).
    content: str = Field(min_length=1, max_length=300_000)
    format: str = Field(pattern=r"^(docx|pdf)$")


def _no_key_stream():
    msg = (
        "**No AI provider configured.**\n\n"
        "Set a free, OpenAI-compatible model in your `.env`, then restart the server:\n"
        "```\nAI_API_KEY=...\nAI_BASE_URL=...\nAI_MODEL=...\n```"
    )
    yield f"data: {json.dumps({'content': msg})}\n\n"
    yield f"data: {json.dumps({'done': True})}\n\n"


def _unknown_type_stream(doc_type: str):
    yield f"data: {json.dumps({'content': f'**Unknown document type:** `{doc_type}`'})}\n\n"
    yield f"data: {json.dumps({'done': True})}\n\n"


# ── GET /api/drafting/types ───────────────────────────────────────────────────
@router.get("/types")
def list_document_types(_: User = Depends(get_current_user)):
    """Return catalogue of available document types and their required fields."""
    return [
        {"id": k, "label": v["label"], "fields": v["fields"]}
        for k, v in DOCUMENT_TYPES.items()
    ]


# ── POST /api/drafting/generate  (SSE streaming) ─────────────────────────────
from app.services.ratelimit import ai_limiter


@router.post("/generate", dependencies=[Depends(ai_limiter)])
def generate_document(
    body: DraftRequest,
    current_user: User = Depends(require_ai_access),   # verification gate (LEGAL-07; flag-gated)
    db: Session = Depends(get_db),
):
    # Entitlements (spec 3.4) — refuse over-quota BEFORE any model work. The paths below
    # that return without calling the model (no key, unknown type, missing field) cost the
    # advocate nothing and are deliberately NOT metered.
    from app.models.billing import KIND_DRAFT
    from app.services.entitlements import enforce_quota, meter
    enforce_quota(db, current_user, KIND_DRAFT)

    from app.ai.llm_config import ai_config
    cfg = ai_config()
    if not cfg["api_key"]:
        return StreamingResponse(_no_key_stream(), media_type="text/event-stream")

    spec = DOCUMENT_TYPES.get(body.document_type)
    if not spec:
        return StreamingResponse(_unknown_type_stream(body.document_type), media_type="text/event-stream")

    # Build user prompt by filling template
    try:
        user_prompt = spec["user_prompt_template"].format(**body.fields)
    except KeyError as exc:
        # Bind the name NOW: Python unbinds `exc` when the except block exits, and this
        # generator runs later (during streaming) — referencing it there raised NameError.
        missing = str(exc).strip("'")

        def _missing_field(field: str = missing):
            yield f"data: {json.dumps({'content': f'**Missing required field:** {field}'})}\n\n"
            yield f"data: {json.dumps({'done': True})}\n\n"
        return StreamingResponse(_missing_field(), media_type="text/event-stream")

    from app.ai.safety import DRAFT_DISCLAIMER
    safety_contract = (
        "\n\nSAFETY CONTRACT (mandatory): This is a DRAFT for advocate review, never a "
        "final document. Never use guarantee/outcome language ('you will win', "
        "'guaranteed', 'file this immediately', 'replaces a lawyer', 'the court will "
        "definitely…'). Leave placeholders like [●] where a fact is unknown rather than "
        "inventing it. Cite only provisions you are confident apply. End the document "
        f"with this exact line:\n\"{DRAFT_DISCLAIMER}\""
    )

    # ── Format reference: authentic Indian drafting skeleton (templates/drafts/) ──
    # Registered types get their own skeleton; a custom document gets the closest
    # keyword match so "any legal document" still lands on a professional structure.
    from app.services import draft_templates
    format_label, skeleton = draft_templates.resolve_format(body.document_type, body.fields)
    format_block = ""
    if skeleton:
        format_block = (
            "\n\nFORMAT REFERENCE (authentic Indian drafting skeleton — follow this "
            "structure and ordering EXACTLY; fill the [●] placeholders from the facts "
            "provided; keep [●] wherever a fact was NOT provided; do not invent names, "
            "dates, numbers or annexures):\n" + "─" * 40 + "\n" + skeleton + "─" * 40
        )

    # ── Statutory grounding: retrieved verbatim provisions the draft may cite ──
    # (same no-hallucination discipline as research: cite ONLY what was retrieved)
    grounding_block = ""
    try:
        from app.ai import rag
        key_facts = " ".join(str(v) for v in list(body.fields.values())[:6])[:200]
        retrieved = rag.retrieve(f"{spec['label']} {key_facts}")
        # Provision-named applications (e.g. "Section 311 CrPC", "Section 349 BNSS"):
        # fetch THAT section's verbatim text deterministically so the application can
        # track the statute's exact language. (CPC Orders/Rules have no verbatim source
        # in the corpus — the prompt instructs the model not to quote them from memory.)
        provision = str(body.fields.get("provision", "")).strip()
        if provision:
            exact = rag.retrieve_by_section(provision)
            if exact:
                retrieved = (exact + "\n\n" + (retrieved or "")).strip()
        if retrieved:
            grounding_block = (
                "\n\nRETRIEVED PROVISIONS (the ONLY statute text you may quote beyond "
                "the sections already named in the format reference; cite exactly as "
                "given):\n" + retrieved
            )
        else:
            grounding_block = (
                "\n\nNOTE: no verified provisions were retrieved for this draft. Cite "
                "only the sections named in the format reference; do not quote statute "
                "text from memory."
            )
    except Exception:
        pass

    # Validation passed and the model is about to be called → record exactly one draft.
    meter(db, current_user, KIND_DRAFT)

    messages = [
        {"role": "system", "content": spec["system_prompt"] + safety_contract + format_block + grounding_block},
        {"role": "user",   "content": user_prompt},
    ]

    async def stream():
        from openai import AsyncOpenAI
        from app.ai.safety import DRAFT_STATUS_REVIEW, DRAFT_DISCLAIMER, AI_GENERATED_NOTICE
        # Generous timeout + retries so a cold local-model load can't surface as an error.
        client = AsyncOpenAI(api_key=cfg["api_key"], base_url=cfg["base_url"],
                             timeout=240.0, max_retries=2)

        # Every draft carries the review status from the first byte (section 2.4)
        yield f"data: {json.dumps({'status': DRAFT_STATUS_REVIEW})}\n\n"
        # Transparency: tell the UI which authentic format skeleton guided this draft
        if format_label:
            yield f"data: {json.dumps({'format_used': format_label})}\n\n"

        full = ""
        try:
            # create(stream=True) — compatible with any OpenAI-compatible provider (free Gemini/Groq/…)
            stream_resp = await client.chat.completions.create(
                model=cfg["model"],
                messages=messages,
                max_tokens=3000,
                temperature=0.2,
                stream=True,
            )
            async for chunk in stream_resp:
                if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                    piece = chunk.choices[0].delta.content
                    full += piece
                    yield f"data: {json.dumps({'content': piece})}\n\n"
            # Always end with the AI-generated disclosure (SC draft AI-in-Courts Regs 2026 reg. 7 / 20(h)).
            # Add the advocate-review disclaimer only if the model didn't already include it (no duplicate).
            tail = chr(10) + chr(10) + '---' + chr(10)
            if DRAFT_DISCLAIMER.split('.')[0].lower() not in full.lower():
                tail += DRAFT_DISCLAIMER + ' '
            tail += AI_GENERATED_NOTICE
            yield f"data: {json.dumps({'content': tail})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        finally:
            yield f"data: {json.dumps({'done': True})}\n\n"

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── Export helpers ──────────────────────────────────────────────────────────────
def _safe_filename(doc_type: str, ext: str) -> str:
    spec = DOCUMENT_TYPES.get(doc_type)
    base = (spec["label"] if spec else doc_type)
    base = re.sub(r"[^A-Za-z0-9]+", "_", base).strip("_").lower() or "document"
    stamp = utcnow().strftime("%Y%m%d")
    return f"{base}_{stamp}.{ext}"


def _build_docx(content: str) -> bytes:
    """Render the generated text into a clean A4 legal-style .docx."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # A4 page with 1-inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Inches(1))

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    lines = content.replace("\r\n", "\n").split("\n")
    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(2)

        # Headings: markdown-style or ALL-CAPS short lines -> bold + centered
        is_md_heading = stripped.startswith("#")
        text = stripped.lstrip("#").strip() if is_md_heading else stripped
        text = text.replace("**", "")
        is_caps_heading = (
            len(text) <= 80
            and text.upper() == text
            and any(c.isalpha() for c in text)
        )

        run = p.add_run(text)
        if is_md_heading or is_caps_heading:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Footer disclaimer
    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(18)
    fr = footer.add_run(
        "Generated by Juriscite. Draft for advocate review. Verify facts, "
        "jurisdiction, limitation, court rules, and latest case law before filing. "
        "AI-assisted content — disclose its AI-generated character if filed in court."
    )
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(content: str) -> bytes:
    """Render the generated text into an A4 legal-style PDF."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import grey
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from xml.sax.saxutils import escape

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=inch, bottomMargin=inch, leftMargin=inch, rightMargin=inch,
        title="Juriscite Document",
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName="Times-Roman",
                          fontSize=12, leading=18, alignment=TA_JUSTIFY)
    head = ParagraphStyle("Head", parent=styles["Normal"], fontName="Times-Bold",
                          fontSize=12, leading=18, alignment=TA_CENTER, spaceAfter=6)
    foot = ParagraphStyle("Foot", parent=styles["Normal"], fontName="Times-Italic",
                          fontSize=8, leading=11, textColor=grey, spaceBefore=18)

    flow = []
    for raw in content.replace("\r\n", "\n").split("\n"):
        stripped = raw.strip()
        if not stripped:
            flow.append(Spacer(1, 8))
            continue
        is_md = stripped.startswith("#")
        text = (stripped.lstrip("#").strip() if is_md else stripped).replace("**", "")
        is_caps = len(text) <= 80 and text.upper() == text and any(c.isalpha() for c in text)
        flow.append(Paragraph(escape(text), head if (is_md or is_caps) else body))

    flow.append(Paragraph(
        "Generated by Juriscite. Draft for advocate review. Verify facts, "
        "jurisdiction, limitation, court rules, and latest case law before filing. "
        "AI-assisted content — disclose its AI-generated character if filed in court.", foot))
    pdf.build(flow)
    return buf.getvalue()


# ── POST /api/drafting/export  (DOCX / PDF download) ────────────────────────────
@router.post("/export")
def export_document(body: ExportRequest, user: User = Depends(get_current_user),
                    db: Session = Depends(get_db)):
    text = (body.content or "").strip()
    if not text:
        return Response(content="No content to export.", status_code=400)

    fmt = body.format.lower()
    try:
        if fmt == "docx":
            data = _build_docx(text)
            media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif fmt == "pdf":
            data = _build_pdf(text)
            media = "application/pdf"
        else:
            return Response(content=f"Unsupported format: {fmt}", status_code=400)
    except ImportError as e:
        return Response(content=f"Export library missing: {e}", status_code=500)
    except Exception as e:
        return Response(content=f"Export failed: {e}", status_code=500)

    filename = _safe_filename(body.document_type, fmt)
    from app.services.tenancy import write_audit
    write_audit(db, tenant_id=user.tenant_id, user_id=user.id,
                action="export_document", entity="DraftExport", detail=f"{body.document_type}.{fmt}")
    return Response(
        content=data,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
